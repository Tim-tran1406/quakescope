"""
Builds the per-session Word study notes for QuakeScope.

Why this exists: the user learns by reading explained, worked examples, and keeps
those explanations as Word docs (personal study notes). This script turns a
session's content into a nicely formatted .docx, with diagrams embedded.

Run it with the project's virtual environment:
    .venv/bin/python tools/build_word_docs.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageChops

# --- paths ---------------------------------------------------------------
ROOT = Path(__file__).resolve().parent.parent
IMAGES = ROOT / "docs" / "images"
OUT_DIR = ROOT / "earthquake projects learning"
OUT_DIR.mkdir(exist_ok=True)


def trim_whitespace(src: Path, dst: Path, pad: int = 20) -> Path:
    """Crop the white border off a PNG so the diagram fills the space."""
    im = Image.open(src).convert("RGB")
    bg = Image.new("RGB", im.size, (255, 255, 255))
    bbox = ImageChops.difference(im, bg).getbbox()
    if bbox:
        left, top, right, bottom = bbox
        left, top = max(0, left - pad), max(0, top - pad)
        right, bottom = min(im.width, right + pad), min(im.height, bottom + pad)
        im = im.crop((left, top, right, bottom))
    im.save(dst)
    return dst


# --- small formatting helpers -------------------------------------------
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


def add_code(doc, text):
    """Add a monospaced code block. Handles multi-line snippets."""
    p = doc.add_paragraph()
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        if i < len(lines) - 1:
            run.add_break()
    return p


# --- build the Session 00 document --------------------------------------
def build_session_00():
    doc = Document()

    doc.add_heading("QuakeScope — Session 00", level=0)
    doc.add_heading("Setup, the stack, and the architecture decision", level=1)
    intro = doc.add_paragraph(
        "These are your study notes for the first session. They explain — in plain "
        "language — what QuakeScope is, the tools we'll learn, how the project is "
        "organised (and why), and the one thing you need to install next. No prior "
        "knowledge is assumed."
    )
    intro.runs[0].italic = True

    # 1. What we're building
    doc.add_heading("1. What we're building", level=2)
    doc.add_paragraph(
        "QuakeScope takes the world's earthquake records and turns them into insight. "
        "The key word is analysis: we don't just plot dots on a map, we answer real "
        "questions:"
    )
    for q in [
        "Is seismic activity really increasing, or are we just detecting more of it?",
        "Which regions are genuinely the most dangerous, and how do we score that?",
        "Do earthquakes cluster into swarms — and can we detect them automatically?",
        "How does the size-vs-frequency pattern (the Gutenberg–Richter law) differ across the globe?",
    ]:
        doc.add_paragraph(q, style="List Bullet")
    doc.add_paragraph(
        "We present the findings two ways: a custom interactive web map (with an API "
        "you build) and a Power BI dashboard."
    )

    # 2. The tools
    doc.add_heading("2. The tools, and why each one exists", level=2)
    doc.add_paragraph(
        "Think of data flowing through a pipeline. Each tool owns one job."
    )
    add_table(
        doc,
        ["Tool", "What it is", "Why it's here"],
        [
            ["Python", "A programming language for working with data", "Fetch, clean, and analyse the earthquake data"],
            ["PostgreSQL", "A database (stores data, answers questions fast)", "Where the data lives"],
            ["SQL", "The language for talking to a database", "Ask questions of the data — the core analyst skill"],
            ["scikit-learn / SciPy", "Python data-science libraries", "Clustering, the b-value, statistics"],
            ["FastAPI", "A tool for building a web API", "A doorway so the website can request our data"],
            ["Leaflet", "A free mapping library", "Draws the interactive earthquake map"],
            ["Power BI", "Microsoft's dashboard tool", "A polished BI dashboard (a top résumé skill)"],
        ],
    )

    # 3. What we set up
    doc.add_heading("3. What we set up today", level=2)
    doc.add_paragraph(
        "A project folder at ~/quakescope, a git repository, and a Python virtual "
        "environment (the .venv folder)."
    )
    p = doc.add_paragraph()
    p.add_run("What's a virtual environment? ").bold = True
    p.add_run(
        "A private box for this project's Python tools, kept separate from every other "
        "project so their library versions never clash."
    )

    # 4. Architecture
    doc.add_heading("4. How the project is organised (investigated 3 ways)", level=2)
    doc.add_paragraph(
        "You asked me to check an architecture decision a few times before committing. "
        "The decision: how should the code and data be laid out? Here's the diagram, "
        "then the three angles I checked."
    )

    trimmed = trim_whitespace(IMAGES / "architecture.png", IMAGES / "architecture_trimmed.png")
    doc.add_picture(str(trimmed), width=Inches(6.2))

    doc.add_heading("Angle 1 — How do professional data teams do it?", level=3)
    doc.add_paragraph(
        "The standard pattern moves data through layers, each cleaner than the last: "
        "raw (exactly as it arrived), staging (cleaned and de-duplicated), and "
        "analytics (the useful end-tables plus our data-science results). The rule of "
        "thumb: never analyse raw data directly — clean it in checkable steps."
    )
    doc.add_heading("Angle 2 — What's easiest for me to learn from?", level=3)
    doc.add_paragraph(
        "One folder per job, named after what it does — no clever abstractions. You "
        "should be able to read the folder names top to bottom and understand the whole "
        "pipeline: get the data → put it in the database → analyse it → serve it → show it."
    )
    doc.add_heading("Angle 3 — What do the two dashboards actually need?", level=3)
    doc.add_paragraph(
        "Both the FastAPI web app and Power BI want the same thing: clean, well-named "
        "tables to read straight from PostgreSQL. So the database is the single source "
        "of truth — we do the hard analysis once, store it, and both dashboards just read it."
    )

    doc.add_heading("The decision (where all three agree)", level=3)
    doc.add_paragraph("Folders in the repo — one per job:")
    add_table(
        doc,
        ["Folder", "Its job"],
        [
            ["src/", "Python code: ingest (get data), load (into DB), analyze (data science)"],
            ["sql/", "SQL that builds the clean tables (raw → staging → analytics)"],
            ["notebooks/", "Scratch space for exploring data before writing real code"],
            ["api/", "The FastAPI web app (added in Phase 4)"],
            ["web/", "The interactive map website (added in Phase 5)"],
            ["docs/", "These learning docs"],
            ["data/raw/", "Downloaded files (not uploaded to GitHub — big and re-downloadable)"],
        ],
    )
    doc.add_paragraph(
        "Inside PostgreSQL we use those same three layers as real separate areas "
        "(Postgres calls them schemas): raw, staging, analytics."
    )
    p = doc.add_paragraph()
    p.add_run("Skipped on purpose: ").bold = True
    p.add_run(
        "a tool called dbt that automates the SQL layer steps. We'll write those steps "
        "by hand first so you understand them, then adding dbt later (optional) will make "
        "sense instead of being magic."
    )

    # 5. Installing PostgreSQL
    doc.add_heading("5. Installing PostgreSQL (done — what we ran and why)", level=2)
    doc.add_paragraph(
        "You already had Homebrew (a package manager — it installs software with one "
        "command), so we used that. Four steps:"
    )
    p = doc.add_paragraph()
    p.add_run("Step 1 — install the database. ").bold = True
    p.add_run("Downloads PostgreSQL 17 and creates the storage area for your data.")
    add_code(doc, "brew install postgresql@17")
    p = doc.add_paragraph()
    p.add_run("Step 2 — let your terminal find it. ").bold = True
    p.add_run("Your terminal only looks in certain folders (its PATH); we add Postgres'.")
    add_code(doc, "echo 'export PATH=\"/opt/homebrew/opt/postgresql@17/bin:$PATH\"' >> ~/.zshrc")
    p = doc.add_paragraph()
    p.add_run("Step 3 — start the server. ").bold = True
    p.add_run("A database must be running to answer questions; this runs it in the background.")
    add_code(doc, "brew services start postgresql@17")
    p = doc.add_paragraph()
    p.add_run("Step 4 — create our project's database. ").bold = True
    p.add_run("PostgreSQL can hold many databases; we made one called quakescope.")
    add_code(doc, "createdb quakescope")
    doc.add_paragraph("Check it worked:")
    add_code(doc, "psql --version\npsql -d quakescope -c \"SELECT version();\"")
    doc.add_paragraph(
        "Your login is 'tim' with no password — normal for local work. The connection "
        "settings are saved in the project's .env file so our Python code can read them."
    )

    # 6. Power BI
    doc.add_heading("6. Heads-up: Power BI on Mac", level=2)
    doc.add_paragraph(
        "Power BI Desktop only runs on Windows. We've parked it until the last phase "
        "(Phase 6), so it never blocks the fun parts. When we get there we'll set up a "
        "free Windows environment together. Everything before that runs natively on your Mac."
    )

    # 7. Next
    doc.add_heading("7. What's next", level=2)
    doc.add_paragraph(
        "Once 'psql --version' works, we start Phase 1: a Python script that pulls real "
        "earthquake data from the USGS. You'll see live data within minutes."
    )

    out = OUT_DIR / "Session 00 - Setup, Stack & Architecture.docx"
    doc.save(str(out))
    return out


def build_session_01():
    doc = Document()

    doc.add_heading("QuakeScope — Session 01", level=0)
    doc.add_heading("Pulling the data (Phase 1)", level=1)
    intro = doc.add_paragraph(
        "What we did: wrote a small Python program that downloaded 124,479 real "
        "earthquakes (2010–2026, magnitude 4.5+, worldwide) from the USGS and saved "
        "them to disk. Below: a picture of how it works, then Python explained from "
        "the simplest building block upward."
    )
    intro.runs[0].italic = True

    trimmed = trim_whitespace(IMAGES / "phase1-ingest.png", IMAGES / "phase1-ingest_trimmed.png")
    doc.add_picture(str(trimmed), width=Inches(5.0))

    doc.add_heading("Part A — Python, from the simplest idea upward", level=2)
    doc.add_paragraph(
        "Each idea is shown with a real line from our script (src/ingest.py). Read top "
        "to bottom — each one is a small step up from the last."
    )
    concepts = [
        ("1. A variable — a name that holds a value",
         "MIN_MAGNITUDE = 4.5",
         "MIN_MAGNITUDE is a label stuck on the number 4.5. Now we write the name "
         "instead of repeating the number, and can change it in one place."),
        ("2. A string — text, written in quotes",
         'USGS_URL = "https://earthquake.usgs.gov/fdsnws/event/1/query"',
         "Anything in quotes is text (a 'string'). This one holds the web address we ask."),
        ("3. An f-string — a value slotted into text",
         'f"{year}-01-01"',
         "The f lets us drop a variable inside text. If year is 2011, this becomes "
         "'2011-01-01'. We build each year's start/end dates this way."),
        ("4. A dictionary — a labelled bag of values",
         'params = {\n    "format": "geojson",\n    "minmagnitude": MIN_MAGNITUDE,\n}',
         "A dictionary stores key -> value pairs: the settings we send to USGS. "
         "Think of it like filling in a form."),
        ("5. Using a library — borrowing code someone else wrote",
         "response = requests.get(USGS_URL, params=params, timeout=180)",
         "requests is a toolbox for talking to websites. .get(...) means 'fetch this "
         "address with these settings'. timeout=180 means 'give up after 3 minutes'."),
        ("6. A function — a named, reusable block of steps",
         "def fetch_one_year(year):\n    ...\n    return response.json()",
         "def defines a little machine: hand it a year, it does the work, and return "
         "hands back the result. Reusable for every year instead of copying code 17 times."),
        ("7. A loop — do the same thing for each item",
         "for year in range(START_YEAR, END_YEAR + 1):\n    ...",
         "range(2010, 2027) is the years 2010-2026. The for loop runs the indented code "
         "once per year — that's how 17 downloads come from a few lines."),
        ("8. An if — make a decision",
         "if out_file.exists():\n    continue",
         "if runs code only when something is true. If this year's file already exists, "
         "continue skips to the next year. That's the 'resume' feature."),
        ("9. try / except — handle things going wrong",
         "try:\n    response = requests.get(...)\nexcept requests.exceptions.RequestException:\n    time.sleep(wait)   # then the loop retries",
         "try runs code that might fail (the network!). If it errors, except catches it "
         "instead of crashing — so we wait and retry. This saved us when 2011 timed out."),
        ("10. Files and JSON — saving the data",
         "out_file.write_text(json.dumps(data))",
         "json.dumps turns Python data into text (JSON is a text format for data); "
         "write_text writes that text into a file on disk."),
    ]
    for title, code, expl in concepts:
        doc.add_heading(title, level=3)
        add_code(doc, code)
        doc.add_paragraph(expl)
    doc.add_paragraph(
        "So the whole script is really a loop (#7) wrapped around a web request (#5) and "
        "saving a file (#10), with #8 and #9 to make it robust."
    )

    doc.add_heading("Part B — Two real-world lessons baked in", level=2)
    p = doc.add_paragraph()
    p.add_run("Validate before building. ").bold = True
    p.add_run(
        "Before writing the script we tested the USGS API live — which caught that "
        "orderby=time-desc is invalid (only time, time-asc, magnitude, magnitude-asc "
        "are allowed) before it became a bug."
    )
    p = doc.add_paragraph()
    p.add_run("Real downloads are flaky, so be resilient. ").bold = True
    p.add_run(
        "Our first run died when 2011 timed out. We added two habits every data engineer "
        "uses: retry (try again after a failure) and resume (skip work already done)."
    )

    doc.add_heading("Part C — Why magnitude 4.5?", level=2)
    doc.add_paragraph(
        "We didn't grab every quake. USGS records every quake of about 4.5+ everywhere "
        "on Earth, but smaller ones are only caught near sensors (mostly the USA). "
        "Including them would make the world look like it mostly shakes in California — "
        "a bias. 4.5+ gives an honest global picture. Questioning whether the data is "
        "fair is the heart of analysis."
    )

    doc.add_heading("Result & what's next", level=2)
    for line in [
        "124,479 earthquakes, 2010–2026, saved as 17 files in data/raw/ (93 MB).",
        "Biggest in our data: M9.1, the 2011 Great Tohoku Earthquake, Japan.",
        "Next (Phase 2): learn SQL — load this into PostgreSQL and shape it into a clean, flat table we can query.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    out = OUT_DIR / "Session 01 - Pulling the Data.docx"
    doc.save(str(out))
    return out


def build_session_02():
    doc = Document()

    doc.add_heading("QuakeScope — Session 02", level=0)
    doc.add_heading("SQL & PostgreSQL (Phase 2)", level=1)
    intro = doc.add_paragraph(
        "What we did: loaded all 124,479 quakes into PostgreSQL, used SQL to turn the "
        "messy JSON into a clean table, then asked the data its first questions."
    )
    intro.runs[0].italic = True

    trimmed = trim_whitespace(IMAGES / "phase2-transform.png", IMAGES / "phase2-transform_trimmed.png")
    doc.add_picture(str(trimmed), width=Inches(6.3))

    doc.add_heading("Part A — SQL, from the simplest idea upward", level=2)
    concepts = [
        ("1. The big picture: database → schema → table",
         "CREATE SCHEMA IF NOT EXISTS staging;",
         "A database (quakescope) is the whole filing cabinet; a schema is a drawer "
         "(we made raw, staging, analytics); a table is a sheet of rows (one per quake) "
         "and columns (the fields)."),
        ("2. Holding raw JSON: the JSONB type",
         "CREATE TABLE raw.events (id text, raw jsonb);",
         "We first copied each quake into raw.events exactly as it arrived, in one "
         "JSONB column (Postgres's JSON format). One quake = one row."),
        ("3. Reading inside JSON, and converting types",
         "(raw -> 'properties' ->> 'mag')::numeric",
         "->  digs into JSON and stays JSON;  ->>  digs in and gives plain text;  "
         "::numeric converts that text into a real number. That is the whole staging "
         "transform: one big SELECT that extracts and types each field."),
        ("4. The simplest query: count the rows",
         "SELECT count(*) FROM staging.events;",
         "SELECT = 'give me'; count(*) = 'how many rows'; FROM = 'from this table'. "
         "Answer: 124,479."),
        ("5. Choose columns, sort, limit: the 5 strongest quakes",
         "SELECT event_time, magnitude, place\nFROM staging.events\nORDER BY magnitude DESC\nLIMIT 5;",
         "List the columns you want; ORDER BY sorts (DESC = biggest first); LIMIT keeps "
         "only the top few."),
        ("6. WHERE — keep only rows that match",
         "SELECT count(*) FROM staging.events\nWHERE magnitude >= 7;",
         "WHERE is the filter: rows that aren't true are dropped before counting. "
         "Result: 252 major quakes."),
        ("7. GROUP BY — bucket rows and summarise each bucket",
         "SELECT date_part('year', event_time)::int AS year,\n       count(*) AS quakes\nFROM staging.events\nGROUP BY year\nORDER BY year;",
         "Collapses 124k rows into one row per year, counting each. Going from raw data "
         "to a summary is the bread and butter of analysis."),
        ("8. Indexes — making questions fast",
         "CREATE INDEX idx_events_time ON staging.events (event_time);",
         "An index is like a book's index: Postgres jumps straight to the rows it needs "
         "instead of scanning all of them. We indexed time, magnitude, and location."),
    ]
    for title, code, expl in concepts:
        doc.add_heading(title, level=3)
        add_code(doc, code)
        doc.add_paragraph(expl)

    doc.add_heading("Part B — What the data already told us", level=2)
    add_table(
        doc,
        ["Question", "SQL idea", "Answer"],
        [
            ["How many quakes?", "count(*)", "124,479"],
            ["Strongest?", "ORDER BY … LIMIT", "M9.1 Tohoku, then M8.8 Kamchatka (2025) & Chile"],
            ["Major (M7+)?", "WHERE", "252"],
            ["Tsunami-flagged?", "WHERE tsunami = 1", "1,381"],
            ["Trend over time?", "GROUP BY year", "flat — ~6,400–9,600/yr, no clear rise"],
        ],
    )
    doc.add_paragraph(
        "That last one is the interesting one: at magnitude 4.5+, the world isn't "
        "obviously shaking more over time. Whether that holds up is a Phase 3 question."
    )

    doc.add_heading("Part C — Two professional habits we used", level=2)
    p = doc.add_paragraph()
    p.add_run("Check data quality. ").bold = True
    p.add_run(
        "We compared count(*) to count(DISTINCT id) — both were 124,479, proving there "
        "are no duplicate quakes. Always sanity-check after loading."
    )
    p = doc.add_paragraph()
    p.add_run("Mind your timezones. ").bold = True
    p.add_run(
        "Per-year counts shifted by a handful versus Phase 1 because Postgres shows "
        "times in your local zone (UTC+7) and a few quakes near New Year cross the "
        "boundary. We'll standardise on UTC in Phase 3."
    )

    doc.add_heading("Result & what's next", level=2)
    for line in [
        "raw.events (124,479 JSONB rows) → staging.events (124,479 clean, typed rows).",
        "You can now answer real questions with SELECT … FROM … WHERE … GROUP BY.",
        "Next (Phase 3): the data science — DBSCAN clusters, the Gutenberg–Richter "
        "b-value, anomalies, and a regional risk score, written into the analytics schema.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    out = OUT_DIR / "Session 02 - SQL & PostgreSQL.docx"
    doc.save(str(out))
    return out


def build_session_03():
    doc = Document()

    doc.add_heading("QuakeScope — Session 03", level=0)
    doc.add_heading("Data Science, Part 1 (Phase 3)", level=1)
    intro = doc.add_paragraph(
        "What we did: two real seismology techniques in Python — found the world's "
        "seismic zones automatically (DBSCAN), and measured the Gutenberg–Richter "
        "b-value — and saved the results into the analytics schema."
    )
    intro.runs[0].italic = True

    # --- Method A ---
    doc.add_heading("Method A — Finding seismic zones with DBSCAN", level=2)
    doc.add_heading("The idea, simplest first", level=3)
    for line in [
        "Clustering = grouping points that are close together.",
        "DBSCAN groups by density: wherever it finds a dense crowd of quakes it calls "
        "that a cluster (a seismic zone); lonely quakes in empty areas are labelled noise.",
        "Settings: a quake within ~50 km of a crowd joins it (eps); a crowd needs at "
        "least 40 quakes to count (min_samples).",
        "We measured distance the round-Earth way (haversine / great-circle), not flat.",
        "The key point: we never told it where the zones are — it discovered them.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_picture(str(IMAGES / "phase3-clusters.png"), width=Inches(6.4))
    doc.add_heading("What it found", level=3)
    doc.add_paragraph(
        "207 seismic zones; 93,675 quakes fell into a zone, 30,804 were scattered noise. "
        "The biggest zones are the Pacific Ring of Fire — and the faint gray trails are "
        "the mid-ocean ridges (too sparse to pass the 40-quake bar)."
    )
    add_table(
        doc,
        ["Zone (top region)", "Quakes", "Strongest"],
        [
            ["Tonga", "11,345", "M8.1"],
            ["Japan", "8,751", "M9.1 (Tohoku)"],
            ["Vanuatu", "8,661", "M8.0"],
            ["Philippines", "7,566", "M7.8"],
            ["Chile", "2,943", "M8.8"],
        ],
    )
    doc.add_paragraph(
        "Why this matters: the result matches real geology with no input from us — the "
        "strongest possible check that the method and data are sound."
    )

    # --- Method B ---
    doc.add_heading("Method B — The Gutenberg–Richter b-value", level=2)
    doc.add_heading("The idea, simplest first", level=3)
    for line in [
        "Small quakes are common; big ones are rare. The Gutenberg–Richter law makes "
        "that precise: each step up in magnitude means a roughly constant drop in count.",
        "On a log scale a constant factor looks like a straight line; its slope is the b-value.",
        "b ≈ 1.0 (the usual worldwide value) means: one whole magnitude up = about 10× fewer quakes.",
        "Mc (magnitude of completeness) is the smallest magnitude we can trust we caught "
        "all of; we only fit the straight line above Mc.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_picture(str(IMAGES / "phase3-bvalue.png"), width=Inches(6.0))
    doc.add_heading("What we got", level=3)
    for line in [
        "Mc = 4.7, b = 1.17, fitted on 70,213 quakes.",
        "The cumulative points lie on a clean straight line — the law holds.",
        "Above ~M7.5 they dip below the line: in only 16 years we haven't had as many "
        "giant quakes as the long-run rate expects (they're rare — need centuries of data).",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("Reading the number", level=3)
    doc.add_paragraph(
        "b near 1.0 means our result is sound. 1.17 is slightly high, most likely because "
        "USGS mixes magnitude types (mb, mww, ms…) — a pro would standardise to one type. "
        "Why anyone cares: a LOW b-value (relatively more big quakes) flags higher hazard, "
        "so comparing b across regions is a real hazard tool."
    )

    doc.add_heading("Saved for the dashboards", level=2)
    for line in [
        "analytics.event_clusters — every quake tagged with its zone id.",
        "analytics.clusters — one row per zone (size, centre, strongest quake, region).",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("What's next — Data Science, Part 2", level=2)
    for line in [
        "Mann–Kendall trend test — settle the 'is activity increasing?' question.",
        "Omori's law — how aftershocks fade after a big quake (e.g. Tohoku).",
        "Isolation Forest + monthly z-scores — find genuinely unusual quakes/periods.",
        "Seismic energy + a regional risk score — rank regions by danger, quantitatively.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    out = OUT_DIR / "Session 03 - Data Science (Part 1).docx"
    doc.save(str(out))
    return out


def build_session_04():
    doc = Document()

    doc.add_heading("QuakeScope — Session 04", level=0)
    doc.add_heading("Data Science, Part 2 (Phase 3)", level=1)
    intro = doc.add_paragraph(
        "What we did: four more real techniques — a trend test, the aftershock-decay "
        "law, two kinds of anomaly detection, and an energy-based risk score — all "
        "saved into the analytics schema."
    )
    intro.runs[0].italic = True

    doc.add_heading("Method C1 — Is seismic activity increasing? (Mann–Kendall)", level=2)
    doc.add_paragraph(
        "Idea, simplest first: a wiggly line of yearly counts fools the eye. The "
        "Mann–Kendall test neutrally asks 'is there a real trend?' and returns a p-value "
        "(below 0.05 = real; above = can't claim one). Sen's slope measures its size."
    )
    doc.add_picture(str(IMAGES / "phase3-trend.png"), width=Inches(6.2))
    doc.add_paragraph(
        "Result: p = 0.93, Sen's slope +3 quakes/year → NO significant trend (2010–2025). "
        "The feeling that earthquakes are increasing is not in the data — it's better news "
        "coverage and more sensors. Being able to say 'I tested it and it's not real' is "
        "exactly the analytical maturity employers want."
    )

    doc.add_heading("Method C2 — Omori's law (how aftershocks fade)", level=2)
    doc.add_paragraph(
        "Idea: after a big quake, aftershocks are frequent then taper off as 1 / time^p "
        "(p ≈ 1). On a log-log chart that's a straight line."
    )
    doc.add_picture(str(IMAGES / "phase3-omori.png"), width=Inches(5.6))
    doc.add_paragraph(
        "Result: the 3,313 aftershocks within 400 km of the 2011 M9.1 Tōhoku quake fade "
        "as t^−0.89 (R² = 0.71) — the law confirmed on your own data."
    )

    doc.add_heading("Method D — Anomaly detection (two views)", level=2)
    p = doc.add_paragraph()
    p.add_run("D1 — Per-quake (Isolation Forest). ").bold = True
    p.add_run(
        "Given each quake's magnitude, depth and location, the algorithm flags weird "
        "combinations. The top hits were all deep-focus quakes:"
    )
    add_table(
        doc,
        ["Quake", "Magnitude", "Depth"],
        [
            ["2013 Sea of Okhotsk", "M8.3", "598 km"],
            ["Bonin Islands, Japan", "M7.8", "664 km"],
            ["Fiji", "M7.9", "671 km"],
        ],
    )
    doc.add_paragraph(
        "That first one is the largest deep-focus earthquake ever recorded — the method "
        "found genuinely special events. Most quakes are shallow; ones 600 km down are rare."
    )
    p = doc.add_paragraph()
    p.add_run("D2 — Per-month (z-score). ").bold = True
    p.add_run("A z-score says how many standard deviations above average a month is; > 2.5 is rare.")
    doc.add_picture(str(IMAGES / "phase3-anomalies-monthly.png"), width=Inches(6.4))
    doc.add_paragraph(
        "The biggest spike is 2011-03 (z = 9.2) — Tōhoku and its aftershocks. Others line "
        "up with real sequences (2021, and 2025-07 = the Kamchatka M8.8)."
    )

    doc.add_heading("Method E — Seismic energy & a regional risk score", level=2)
    doc.add_paragraph(
        "Idea: counting quakes treats a M5 and a M9 the same, which is absurd — energy ≈ "
        "10^(1.5 × magnitude + 4.8) joules, so each +1 magnitude ≈ 32× more energy. We add "
        "energy per zone and blend three danger factors into a 0–100 score: total energy "
        "(50%) + strongest quake (30%) + how often it shakes (20%)."
    )
    doc.add_picture(str(IMAGES / "phase3-risk.png"), width=Inches(6.2))
    for line in [
        "Tōhoku alone released ~30% of ALL the seismic energy in 16 years.",
        "The 15 quakes of M8+ (0.01% of events) released ~70% of all energy — a handful "
        "of giants dominate, which is why energy (not counts) is the honest measure.",
        "Riskiest zones: Japan (95), Tonga, Kamchatka/Russia, Chile, Vanuatu.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("The analytics schema is now full", level=2)
    doc.add_paragraph(
        "Four results tables, ready for the API and dashboards: event_clusters, clusters, "
        "event_anomalies, zone_risk."
    )
    doc.add_heading("What's next — Phase 4: FastAPI", level=2)
    doc.add_paragraph(
        "We've done the thinking; now we serve it. We'll build a small web API so other "
        "programs (the map website, and anyone online) can ask for this data."
    )

    out = OUT_DIR / "Session 04 - Data Science (Part 2).docx"
    doc.save(str(out))
    return out


def build_session_05():
    doc = Document()

    doc.add_heading("QuakeScope — Session 05", level=0)
    doc.add_heading("FastAPI (Phase 4)", level=1)
    intro = doc.add_paragraph(
        "What we did: built a small web API — a set of URLs that hand back our earthquake "
        "data and analysis as JSON, read live from PostgreSQL. This is the backend the "
        "Phase 5 map website will talk to."
    )
    intro.runs[0].italic = True

    doc.add_picture(str(IMAGES / "phase4-api.png"), width=Inches(6.4))

    doc.add_heading("Part A — FastAPI, from the simplest idea upward", level=2)
    concepts = [
        ("1. What is an API? An endpoint?",
         None,
         "An API is a doorway for programs; an endpoint is one URL behind it. When you "
         "'GET' a URL like /zones/risk, the API runs some code and hands back data as JSON."),
        ("2. Your first endpoint: a function becomes a URL",
         '@app.get("/")\ndef root():\n    return {"name": "QuakeScope API"}',
         "@app.get('/') says 'when someone visits /, run this function'. Whatever you "
         "return (here a dictionary) FastAPI turns into JSON automatically."),
        ("3. Path parameters — part of the URL",
         '@app.get("/quakes/{quake_id}")\ndef get_quake(quake_id: str):',
         "{quake_id} is a blank in the URL. Visiting /quakes/us6000qw60 passes that id "
         "straight into the function."),
        ("4. Query parameters — the bits after ?, with built-in checks",
         "min_magnitude: float = Query(4.5, ge=0, le=10)",
         "/quakes?min_magnitude=7 fills this in. The ge/le rules make FastAPI reject bad "
         "input (e.g. magnitude 99) before our code even runs."),
        ("5. Reading the database — safely",
         'conn.execute(text("... WHERE magnitude >= :minmag"), {"minmag": value})',
         "We never glue user input directly into SQL (the classic 'SQL injection' hole). "
         "We leave a labelled blank :minmag and pass the value separately, so the database "
         "treats it as data, never as code."),
        ("6. Rows → JSON, and errors",
         "raise HTTPException(404, 'earthquake not found')",
         "A list of rows becomes a JSON array automatically (dates and numbers included). "
         "If something isn't found we raise a proper 404 reply."),
        ("7. CORS — letting a website call us",
         "app.add_middleware(CORSMiddleware, allow_origins=['*'])",
         "One line gives a browser permission to call this API from a different address — "
         "needed so the Phase 5 map page can fetch data."),
        ("8. The free gift: interactive docs",
         None,
         "FastAPI reads our function signatures and builds a clickable test page at /docs "
         "— no work from us. You can try every endpoint in the browser."),
        ("9. Running it",
         ".venv/bin/uvicorn api.main:app --reload",
         "uvicorn is the server that runs the app; --reload restarts it whenever you save."),
    ]
    for title, code, expl in concepts:
        doc.add_heading(title, level=3)
        if code:
            add_code(doc, code)
        doc.add_paragraph(expl)

    doc.add_heading("Part B — The endpoints we built", level=2)
    add_table(
        doc,
        ["Endpoint", "Returns", "Reads from"],
        [
            ["/quakes", "quakes, filterable by magnitude/date/region", "staging.events"],
            ["/quakes/{id}", "one quake", "staging.events"],
            ["/stats/yearly", "quakes per year", "staging.events"],
            ["/zones/risk", "riskiest zones (Method E)", "analytics.zone_risk"],
            ["/clusters", "seismic zones (Method A)", "analytics.clusters"],
            ["/anomalies", "most unusual quakes (Method D)", "analytics.event_anomalies + staging.events"],
        ],
    )
    doc.add_paragraph(
        "We tested them live — /zones/risk returned Japan (95.4), Tonga, Kamchatka, Chile, "
        "Vanuatu; /anomalies returned the Sea of Okhotsk deep quakes. Each reads straight "
        "from the database."
    )

    doc.add_heading("Try it yourself", level=2)
    add_code(doc, "cd ~/quakescope\n.venv/bin/uvicorn api.main:app --reload")
    doc.add_paragraph("Then open http://127.0.0.1:8000/docs in your browser and click around.")

    doc.add_heading("What's next — Phase 5: the web map", level=2)
    doc.add_paragraph(
        "Now that the data is served, we'll build an interactive map website (Leaflet) "
        "that calls these endpoints and plots the quakes, zones and risk live in your "
        "browser — your own custom dashboard."
    )

    out = OUT_DIR / "Session 05 - FastAPI.docx"
    doc.save(str(out))
    return out


def build_session_06():
    doc = Document()

    doc.add_heading("QuakeScope — Session 06", level=0)
    doc.add_heading("The interactive web map (Phase 5)", level=1)
    intro = doc.add_paragraph(
        "What we did: built a map website that calls your API and plots every quake and "
        "risk zone, live, in the browser — the first piece you can actually show people."
    )
    intro.runs[0].italic = True

    doc.add_picture(str(IMAGES / "phase5-map.png"), width=Inches(6.5))
    doc.add_paragraph(
        "Red = shallow, amber = medium, blue = deep; dot size = magnitude; orange circles "
        "= risk zones. The shape is the Pacific Ring of Fire — drawn from live API data."
    ).runs[0].italic = True

    doc.add_heading("How a web page works, from the simplest idea upward", level=2)
    doc.add_paragraph(
        "A website is three plain-text languages working together — we used all three."
    )
    concepts = [
        ("1. HTML — the structure (the bones)",
         '<div id="map"></div>',
         "HTML describes what's on the page using tags. Ours is tiny: a header bar, an "
         "empty box for the map, and the legend."),
        ("2. CSS — the styling (the looks)",
         "#map { height: 100vh; width: 100%; }",
         "CSS says how things look — colours, sizes, positions. 100vh means full screen height."),
        ("3. A library from a CDN — borrowing the map engine",
         '<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>',
         "We didn't build a map from scratch — we loaded Leaflet (a free map library) from "
         "its public address (a 'CDN')."),
        ("4. Making the map (JavaScript)",
         'const map = L.map("map").setView([10, 150], 3);',
         "JavaScript is the behaviour — code that runs in the browser. This centres the map "
         "on the Pacific and adds the background map tiles."),
        ("5. fetch() — calling YOUR API from the browser",
         "const res = await fetch(`/quakes?min_magnitude=${minMag}&limit=5000`);\nconst quakes = await res.json();",
         "fetch asks the API for data; await means 'wait for the reply'; .json() turns the "
         "reply into usable data — the same JSON you saw in Phase 4."),
        ("6. Drawing the data",
         'L.circleMarker([q.latitude, q.longitude], {radius: magRadius(q.magnitude),\n  fillColor: depthColour(q.depth_km)}).bindPopup(`M${q.magnitude}`).addTo(quakeLayer);',
         "For each quake we drop a coloured, sized dot with a click-popup."),
        ("7. Layers & interactivity",
         None,
         "Two layers (quakes, risk zones) with a toggle, and M5/M6/M7 buttons that re-run "
         "fetch with a new magnitude filter — so the map updates instantly when you click."),
    ]
    for title, code, expl in concepts:
        doc.add_heading(title, level=3)
        if code:
            add_code(doc, code)
        doc.add_paragraph(expl)

    doc.add_heading("How it all connects", level=2)
    doc.add_paragraph(
        "The map is served by FastAPI at /app, so it lives at the same address as the API. "
        "That's why it can call /quakes and /zones/risk directly — one server doing both "
        "jobs, no CORS headaches."
    )

    doc.add_heading("Run it yourself", level=2)
    add_code(doc, "cd ~/quakescope\n.venv/bin/uvicorn api.main:app --reload")
    doc.add_paragraph(
        "Then open http://127.0.0.1:8000/app/ — pan, zoom, click a quake, toggle layers, "
        "switch magnitudes."
    )

    doc.add_heading("What's next — Phase 6: Power BI", level=2)
    doc.add_paragraph(
        "The last build phase, and the résumé skill we saved for the end. We'll set up a "
        "free Windows environment (Power BI Desktop is Windows-only), connect Power BI "
        "straight to PostgreSQL, and build a polished BI dashboard from the analytics tables."
    )

    out = OUT_DIR / "Session 06 - Web Map.docx"
    doc.save(str(out))
    return out


def build_session_07():
    doc = Document()

    doc.add_heading("QuakeScope — Session 07", level=0)
    doc.add_heading("Polish & reproducibility (Phase 7)", level=1)
    intro = doc.add_paragraph(
        "What we did: turned a working project into a presentable, reproducible one — the "
        "part that makes a portfolio project look professional rather than a pile of scripts."
    )
    intro.runs[0].italic = True

    doc.add_picture(str(IMAGES / "architecture.png"), width=Inches(6.3))

    doc.add_heading("Why this phase matters", level=2)
    doc.add_paragraph(
        "A recruiter or teammate spends ~30 seconds deciding if your project is any good. "
        "They look at two things: the README, and whether they can run it. Phase 7 nails both."
    )

    doc.add_heading("1. The README — your project's front door", level=3)
    doc.add_paragraph(
        "README.md is the first (often only) thing people read. Ours leads with a screenshot, "
        "states what the project does, lists the findings (the actual insights — the part that "
        "stands out), shows the stack and architecture, and gives copy-paste run steps. Tip: "
        "lead with results, not setup — 'earthquakes aren't increasing (p = 0.93)' beats "
        "'this uses Python and PostgreSQL'."
    )

    doc.add_heading("2. Reproducibility — pinned dependencies", level=3)
    add_code(doc, "pandas==3.0.3\nscikit-learn==1.9.0")
    doc.add_paragraph(
        "We pinned exact versions in requirements.txt. Without ==, someone installing next "
        "year could get newer libraries that behave differently and break things. Pinning lets "
        "anyone recreate the exact environment that produced these results. Doc/screenshot-only "
        "tools live in requirements-dev.txt so the core project stays lean."
    )

    doc.add_heading("3. One command to rebuild everything", level=3)
    add_code(doc, "./run_all.sh")
    doc.add_paragraph(
        "This runs the whole pipeline in order — ingest → load → SQL → all six analyses — so the "
        "project can be regenerated from nothing. Two habits made it possible: idempotent steps "
        "(ingest resumes; load_raw drops & recreates) so it's safe to re-run, and one script per "
        "job so the order is obvious."
    )

    doc.add_heading("4. A clean repository (.gitignore)", level=3)
    for line in [
        "data/raw/ — 93 MB of re-downloadable data (code, not data, belongs in git).",
        ".venv/ — the installed libraries (rebuilt from requirements.txt).",
        ".env — database secrets (never commit credentials).",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph(
        "So what's on GitHub is exactly what matters: the code, the SQL, the docs, the README."
    )

    doc.add_heading("Optional next step — free cloud deploy", level=2)
    doc.add_paragraph(
        "To put it online for free: host PostgreSQL on a free tier (e.g. Supabase), deploy the "
        "FastAPI app on a free host (e.g. Render or Railway), and point the map at that URL. "
        "Left as an optional follow-up — it needs accounts and secrets, which are yours to create."
    )

    doc.add_heading("What's left", level=2)
    doc.add_paragraph(
        "Phase 6 — Power BI (deferred): connect Power BI to this same PostgreSQL database and "
        "build a BI dashboard, once a Windows environment is set up. Otherwise the project is "
        "complete and presentable: ingestion → database → data science → API → interactive map, "
        "reproducible with one command."
    )

    out = OUT_DIR / "Session 07 - Polish & Reproducibility.docx"
    doc.save(str(out))
    return out


def build_session_08():
    doc = Document()

    doc.add_heading("QuakeScope — Session 08", level=0)
    doc.add_heading("How to make a website", level=1)
    intro = doc.add_paragraph(
        "A plain-language guide to how a website is built and put online, using the "
        "QuakeScope map as the worked example. No prior web knowledge assumed."
    )
    intro.runs[0].italic = True

    doc.add_heading("1. What a website actually is", level=2)
    doc.add_paragraph("A website is just files a browser knows how to read. Three languages, three jobs:")
    for line in [
        "HTML is the structure (the bones): what is on the page.",
        "CSS is the style (the looks): colours, spacing, fonts.",
        "JavaScript is the behaviour (the actions): it runs in the browser and makes things happen.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. Two kinds of website: static and dynamic", level=2)
    doc.add_picture(str(IMAGES / "static-vs-api.png"), width=Inches(6.4))
    doc.add_paragraph(
        "A static website is finished files handed over as-is, with no server thinking — "
        "cheap, fast, hard to break. A dynamic website has a server (often a database and "
        "an API) that builds a fresh answer per visit. Rule of thumb: if the data does not "
        "change per visitor, make it static. Earthquake history is fixed, so our map is "
        "static — we export the data to JSON once and the browser does the rest."
    )

    doc.add_heading("3. The pieces of a simple website", level=2)
    for line in [
        "index.html — the page itself (the browser looks for this first).",
        "A stylesheet — ours is inline; bigger sites use a separate .css file.",
        "JavaScript — app.js, the behaviour.",
        "Data and libraries — our data/quakes.json, plus Leaflet loaded from a public CDN.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("4. How putting it online works", level=2)
    doc.add_paragraph(
        "A website lives on a host: a computer, always on, that hands your files to anyone "
        "who asks. You do not need your own server. The simplest free host for a code "
        "project is GitHub Pages, which serves files straight from your repository at an "
        "address like https://username.github.io/project/."
    )

    doc.add_heading("5. How we shipped the QuakeScope map", level=2)
    for line in [
        "Built the page — web/index.html and web/app.js (using Leaflet for the map).",
        "Exported the data — a Python script wrote web/data/quakes.json, so no backend is needed.",
        "Tested locally — ran a tiny local web server (you can't just double-click the HTML, "
        "because browsers block a page from loading data files straight off the disk).",
        "Deployed to GitHub Pages — an automated job publishes the web/ folder on every push.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("6. Security: what to watch on any website", level=2)
    for line in [
        "Anything sent to the browser is public — never put passwords, private keys, or personal data in HTML/CSS/JS/data.",
        "Keep secrets in .env and out of git, so they never reach GitHub or the site.",
        "Check what you publish — we scanned for passwords and internal addresses before going live.",
        "A static site has almost no attack surface: no server, no exposed database, nothing to hack.",
        "If you later add a real backend, validate input, never glue user text into SQL, and lock down who can call your API.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("The result", level=2)
    doc.add_paragraph("The live map, hosted free as a static site:")
    doc.add_picture(str(IMAGES / "phase5-map.png"), width=Inches(6.4))

    out = OUT_DIR / "Session 08 - How to Make a Website.docx"
    doc.save(str(out))
    return out


if __name__ == "__main__":
    for builder in (build_session_00, build_session_01, build_session_02, build_session_03,
                    build_session_04, build_session_05, build_session_06, build_session_07,
                    build_session_08):
        print(f"Wrote: {builder()}")
